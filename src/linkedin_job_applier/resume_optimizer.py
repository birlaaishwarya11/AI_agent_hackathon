"""AI-powered resume optimization"""

import logging
import json
import re
from typing import Dict, List, Any, Optional
from datetime import datetime, timedelta
from pathlib import Path

import openai
import anthropic
from docx import Document
from docx.shared import Inches

from .config import Config
from .resume_analyzer import ResumeAnalyzer
from .job_matcher import JobMatcher

logger = logging.getLogger(__name__)

class ResumeOptimizer:
    """AI-powered resume optimization for specific job applications"""
    
    def __init__(self, config: Config):
        self.config = config
        self.resume_analyzer = ResumeAnalyzer(config)
        self.job_matcher = JobMatcher(config)
        
        # Initialize AI clients
        ai_config = self.config.get_ai_config()
        
        if ai_config["openai_api_key"]:
            openai.api_key = ai_config["openai_api_key"]
            self.openai_client = openai.OpenAI(api_key=ai_config["openai_api_key"])
        else:
            self.openai_client = None
        
        if ai_config["anthropic_api_key"]:
            self.anthropic_client = anthropic.Anthropic(api_key=ai_config["anthropic_api_key"])
        else:
            self.anthropic_client = None
        
        self.preferred_model = ai_config["preferred_model"]
    
    async def optimize_resume(self, job_id: str, job_description: str, 
                            optimization_level: str = "moderate") -> Dict[str, Any]:
        """Optimize resume for a specific job"""
        try:
            # Get resume data
            resume_data = await self.resume_analyzer.get_resume_data()
            if not resume_data:
                return {"error": "No resume analysis available. Please analyze resume first."}
            
            # Get job match analysis
            match_result = await self.job_matcher.match_job_resume(job_id, job_description)
            if "error" in match_result:
                return match_result
            
            # Generate optimization strategy
            optimization_strategy = await self._create_optimization_strategy(
                resume_data["analysis"],
                match_result["match_result"],
                job_description,
                optimization_level
            )
            
            # Generate optimized resume content
            optimized_content = await self._generate_optimized_content(
                resume_data,
                optimization_strategy,
                job_description
            )
            
            # Create optimized resume document
            optimized_resume_path = await self._create_optimized_resume_document(
                job_id,
                optimized_content,
                optimization_level
            )
            
            # Save optimization record
            optimization_record = {
                "job_id": job_id,
                "optimization_level": optimization_level,
                "original_resume": resume_data["file_path"],
                "optimized_resume": str(optimized_resume_path),
                "optimization_strategy": optimization_strategy,
                "match_improvement": await self._calculate_match_improvement(
                    match_result["match_result"]["overall_score"],
                    optimized_content
                ),
                "optimized_at": datetime.now().isoformat()
            }
            
            await self._save_optimization_record(job_id, optimization_record)
            
            return {
                "success": True,
                "optimized_resume_path": str(optimized_resume_path),
                "optimization_record": optimization_record
            }
            
        except Exception as e:
            logger.error(f"Error optimizing resume: {e}")
            return {"error": str(e)}
    
    async def _create_optimization_strategy(self, resume_analysis: Dict[str, Any],
                                          match_result: Dict[str, Any],
                                          job_description: str,
                                          optimization_level: str) -> Dict[str, Any]:
        """Create optimization strategy based on job requirements and match analysis"""
        
        strategy = {
            "focus_areas": [],
            "keywords_to_add": [],
            "skills_to_emphasize": [],
            "sections_to_modify": [],
            "content_adjustments": [],
            "optimization_level": optimization_level
        }
        
        # Analyze gaps and opportunities
        missing_skills = match_result["requirements_missing"]
        met_requirements = match_result["requirements_met"]
        detailed_analysis = match_result["detailed_analysis"]
        
        # Determine focus areas based on match scores
        category_scores = match_result["category_scores"]
        
        if category_scores["technical_skills"] < 0.7:
            strategy["focus_areas"].append("technical_skills")
        
        if category_scores["keywords"] < 0.6:
            strategy["focus_areas"].append("keyword_optimization")
        
        if category_scores["experience"] < 0.7:
            strategy["focus_areas"].append("experience_presentation")
        
        # Keywords to add (from missing requirements)
        strategy["keywords_to_add"] = (
            missing_skills.get("technical_skills", [])[:5] +  # Top 5 missing tech skills
            missing_skills.get("must_have", [])[:3]  # Top 3 missing must-haves
        )
        
        # Skills to emphasize (from met requirements)
        strategy["skills_to_emphasize"] = (
            met_requirements.get("technical_skills", []) +
            met_requirements.get("nice_to_have", [])
        )
        
        # Determine sections to modify based on optimization level
        if optimization_level == "light":
            strategy["sections_to_modify"] = ["summary", "skills"]
        elif optimization_level == "moderate":
            strategy["sections_to_modify"] = ["summary", "skills", "experience"]
        else:  # aggressive
            strategy["sections_to_modify"] = ["summary", "skills", "experience", "projects", "education"]
        
        # Content adjustment recommendations
        if "technical_skills" in strategy["focus_areas"]:
            strategy["content_adjustments"].append("Enhance technical skills section with job-relevant technologies")
        
        if "keyword_optimization" in strategy["focus_areas"]:
            strategy["content_adjustments"].append("Integrate job-specific keywords naturally throughout resume")
        
        if "experience_presentation" in strategy["focus_areas"]:
            strategy["content_adjustments"].append("Reframe experience to highlight relevant accomplishments")
        
        return strategy
    
    async def _generate_optimized_content(self, resume_data: Dict[str, Any],
                                        optimization_strategy: Dict[str, Any],
                                        job_description: str) -> Dict[str, Any]:
        """Generate optimized resume content using AI"""
        
        # Prepare prompt for AI
        prompt = self._create_optimization_prompt(
            resume_data,
            optimization_strategy,
            job_description
        )
        
        # Generate optimized content using preferred AI model
        if self.preferred_model.startswith("gpt") and self.openai_client:
            optimized_content = await self._generate_with_openai(prompt)
        elif self.preferred_model.startswith("claude") and self.anthropic_client:
            optimized_content = await self._generate_with_anthropic(prompt)
        else:
            raise ValueError("No AI client available for resume optimization")
        
        return optimized_content
    
    def _create_optimization_prompt(self, resume_data: Dict[str, Any],
                                  optimization_strategy: Dict[str, Any],
                                  job_description: str) -> str:
        """Create prompt for AI resume optimization"""
        
        original_text = resume_data["text"]
        analysis = resume_data["analysis"]
        strategy = optimization_strategy
        
        prompt = f"""
You are an expert resume writer and career coach. Your task is to optimize a resume for a specific job application.

ORIGINAL RESUME TEXT:
{original_text}

CURRENT RESUME ANALYSIS:
- Technical Skills: {', '.join(analysis.get('technical_skills', []))}
- Soft Skills: {', '.join(analysis.get('soft_skills', []))}
- Keywords: {', '.join(analysis.get('keywords', []))}
- Experience Years: {analysis.get('experience_years', 'Not specified')}

JOB DESCRIPTION:
{job_description}

OPTIMIZATION STRATEGY:
- Focus Areas: {', '.join(strategy['focus_areas'])}
- Keywords to Add: {', '.join(strategy['keywords_to_add'])}
- Skills to Emphasize: {', '.join(strategy['skills_to_emphasize'])}
- Sections to Modify: {', '.join(strategy['sections_to_modify'])}
- Optimization Level: {strategy['optimization_level']}

INSTRUCTIONS:
1. Optimize the resume to better match the job requirements
2. Naturally integrate the missing keywords and skills
3. Emphasize relevant experience and accomplishments
4. Maintain truthfulness - do not add false information
5. Keep the same overall structure and format
6. Focus on the specified sections for modification
7. Ensure the optimized resume flows naturally and professionally

Please provide the optimized resume content in the following JSON format:
{{
    "summary": "Optimized professional summary",
    "technical_skills": ["list", "of", "technical", "skills"],
    "soft_skills": ["list", "of", "soft", "skills"],
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "duration": "Start - End",
            "description": "Optimized job description with relevant keywords and accomplishments"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "Institution Name",
            "year": "Year",
            "details": "Relevant details"
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Project description with relevant technologies",
            "technologies": ["tech1", "tech2"]
        }}
    ],
    "certifications": ["list", "of", "certifications"],
    "optimization_notes": "Brief notes on what was optimized and why"
}}
"""
        
        return prompt
    
    async def _generate_with_openai(self, prompt: str) -> Dict[str, Any]:
        """Generate optimized content using OpenAI"""
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer and career coach."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=3000
            )
            
            content = response.choices[0].message.content
            
            # Parse JSON response
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                raise ValueError("Could not parse AI response as JSON")
                
        except Exception as e:
            logger.error(f"Error generating content with OpenAI: {e}")
            raise
    
    async def _generate_with_anthropic(self, prompt: str) -> Dict[str, Any]:
        """Generate optimized content using Anthropic Claude"""
        try:
            response = self.anthropic_client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=3000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            content = response.content[0].text
            
            # Parse JSON response
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                raise ValueError("Could not parse AI response as JSON")
                
        except Exception as e:
            logger.error(f"Error generating content with Anthropic: {e}")
            raise
    
    async def _create_optimized_resume_document(self, job_id: str, 
                                              optimized_content: Dict[str, Any],
                                              optimization_level: str) -> Path:
        """Create optimized resume document"""
        try:
            # Create filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_optimized_{job_id}_{optimization_level}_{timestamp}.docx"
            output_path = self.config.data_dir / "optimized_resumes" / filename
            
            # Create directory if it doesn't exist
            output_path.parent.mkdir(exist_ok=True)
            
            # Create new document
            doc = Document()
            
            # Add header with contact info (if available)
            resume_data = await self.resume_analyzer.get_resume_data()
            if resume_data:
                contact_info = resume_data["analysis"].get("contact_info", {})
                if contact_info.get("email"):
                    header = doc.sections[0].header
                    header_para = header.paragraphs[0]
                    header_para.text = f"Email: {contact_info['email']}"
                    if contact_info.get("phone"):
                        header_para.text += f" | Phone: {contact_info['phone']}"
                    if contact_info.get("linkedin"):
                        header_para.text += f" | LinkedIn: {contact_info['linkedin']}"
            
            # Add professional summary
            if optimized_content.get("summary"):
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(optimized_content["summary"])
            
            # Add technical skills
            if optimized_content.get("technical_skills"):
                doc.add_heading("Technical Skills", level=1)
                skills_text = ", ".join(optimized_content["technical_skills"])
                doc.add_paragraph(skills_text)
            
            # Add experience
            if optimized_content.get("experience"):
                doc.add_heading("Professional Experience", level=1)
                for exp in optimized_content["experience"]:
                    # Job title and company
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(f"{exp['title']} - {exp['company']}")
                    title_run.bold = True
                    
                    # Duration
                    if exp.get("duration"):
                        duration_para = doc.add_paragraph(exp["duration"])
                        duration_para.style = "Intense Quote"
                    
                    # Description
                    if exp.get("description"):
                        doc.add_paragraph(exp["description"])
                    
                    doc.add_paragraph()  # Add space
            
            # Add projects
            if optimized_content.get("projects"):
                doc.add_heading("Projects", level=1)
                for project in optimized_content["projects"]:
                    # Project name
                    project_para = doc.add_paragraph()
                    project_run = project_para.add_run(project["name"])
                    project_run.bold = True
                    
                    # Description
                    if project.get("description"):
                        doc.add_paragraph(project["description"])
                    
                    # Technologies
                    if project.get("technologies"):
                        tech_text = "Technologies: " + ", ".join(project["technologies"])
                        tech_para = doc.add_paragraph(tech_text)
                        tech_para.style = "Intense Quote"
                    
                    doc.add_paragraph()  # Add space
            
            # Add education
            if optimized_content.get("education"):
                doc.add_heading("Education", level=1)
                for edu in optimized_content["education"]:
                    # Degree and institution
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{edu['degree']} - {edu['institution']}")
                    edu_run.bold = True
                    
                    # Year
                    if edu.get("year"):
                        year_para = doc.add_paragraph(edu["year"])
                        year_para.style = "Intense Quote"
                    
                    # Details
                    if edu.get("details"):
                        doc.add_paragraph(edu["details"])
            
            # Add certifications
            if optimized_content.get("certifications"):
                doc.add_heading("Certifications", level=1)
                cert_text = "\n".join([f"• {cert}" for cert in optimized_content["certifications"]])
                doc.add_paragraph(cert_text)
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Optimized resume saved to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating optimized resume document: {e}")
            raise
    
    async def _calculate_match_improvement(self, original_score: float, 
                                         optimized_content: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate estimated match improvement"""
        try:
            # This is a simplified estimation - in practice, you'd want to
            # run the optimized content through the job matcher
            
            # Count added keywords and skills
            added_keywords = len(optimized_content.get("technical_skills", []))
            added_soft_skills = len(optimized_content.get("soft_skills", []))
            
            # Estimate improvement based on additions
            keyword_boost = min(0.1, added_keywords * 0.01)
            skill_boost = min(0.05, added_soft_skills * 0.01)
            
            estimated_new_score = min(1.0, original_score + keyword_boost + skill_boost)
            improvement = estimated_new_score - original_score
            
            return {
                "original_score": original_score,
                "estimated_new_score": round(estimated_new_score, 3),
                "improvement": round(improvement, 3),
                "improvement_percentage": round(improvement * 100, 1)
            }
            
        except Exception as e:
            logger.error(f"Error calculating match improvement: {e}")
            return {
                "original_score": original_score,
                "estimated_new_score": original_score,
                "improvement": 0.0,
                "improvement_percentage": 0.0
            }
    
    async def _save_optimization_record(self, job_id: str, 
                                      optimization_record: Dict[str, Any]) -> None:
        """Save optimization record to file"""
        try:
            records_file = self.config.data_dir / "optimization_records.json"
            
            # Load existing records
            records = {}
            if records_file.exists():
                with open(records_file, 'r') as f:
                    records = json.load(f)
            
            # Add new record
            records[job_id] = optimization_record
            
            # Save updated records
            with open(records_file, 'w') as f:
                json.dump(records, f, indent=2)
                
        except Exception as e:
            logger.error(f"Error saving optimization record: {e}")
    
    async def get_optimization_history(self, days_back: int = 30) -> List[Dict[str, Any]]:
        """Get optimization history"""
        try:
            records_file = self.config.data_dir / "optimization_records.json"
            
            if not records_file.exists():
                return []
            
            with open(records_file, 'r') as f:
                records = json.load(f)
            
            # Convert to list and filter by date if needed
            history = []
            cutoff_date = datetime.now() - timedelta(days=days_back)
            
            for job_id, record in records.items():
                optimized_date = datetime.fromisoformat(record["optimized_at"])
                if optimized_date >= cutoff_date:
                    history.append({
                        "job_id": job_id,
                        **record
                    })
            
            # Sort by optimization date (newest first)
            history.sort(key=lambda x: x["optimized_at"], reverse=True)
            
            return history
            
        except Exception as e:
            logger.error(f"Error getting optimization history: {e}")
            return []